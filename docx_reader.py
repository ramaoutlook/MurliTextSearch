"""
Reads a .docx murli file into:
  - a single whitespace-normalized text stream (for searching/sentence splitting)
  - a list of paragraph boundaries (start offset in the stream -> paragraph index)
  - a list of date headings found (offset in the stream -> date string)
"""
import re
from dataclasses import dataclass
from typing import List

import docx

from date_detector import is_date_heading, extract_date_string


def _normalize(text: str) -> str:
    """Collapse all whitespace (including line-wrap newlines/tabs) to single spaces."""
    return re.sub(r'\s+', ' ', text).strip()


@dataclass
class DateHeading:
    offset: int          # character offset in the normalized stream where this heading starts
    date_string: str      # e.g. "21-01-69" or "१-१-९४"
    paragraph_index: int


@dataclass
class ParsedDocument:
    stream: str                      # full normalized text of the whole document
    paragraph_offsets: List[int]      # start offset of each paragraph in `stream`
    date_headings: List[DateHeading]  # in document order
    murli_type: str
    source_path: str


def parse_docx(file_path: str, murli_type: str) -> ParsedDocument:
    document = docx.Document(file_path)

    stream_parts: List[str] = []
    paragraph_offsets: List[int] = []
    date_headings: List[DateHeading] = []

    cursor = 0
    for idx, paragraph in enumerate(document.paragraphs):
        raw_text = paragraph.text
        norm_text = _normalize(raw_text)

        paragraph_offsets.append(cursor)

        if norm_text and is_date_heading(norm_text, murli_type):
            date_headings.append(
                DateHeading(
                    offset=cursor,
                    date_string=extract_date_string(norm_text, murli_type),
                    paragraph_index=idx,
                )
            )

        if norm_text:
            stream_parts.append(norm_text)
            cursor += len(norm_text)
            # separate paragraphs with a single space in the joined stream
            stream_parts.append(' ')
            cursor += 1

    stream = ''.join(stream_parts)

    return ParsedDocument(
        stream=stream,
        paragraph_offsets=paragraph_offsets,
        date_headings=date_headings,
        murli_type=murli_type,
        source_path=file_path,
    )
