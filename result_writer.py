"""
Writes the final result .docx: one document per API call, all matches listed
in the order they were found, each as its own paragraph/block showing the
date, page number, and the 5-sentence excerpt with the matched phrase bolded.
"""
import os
import re
import unicodedata
from datetime import datetime
from typing import List

import docx
from docx.shared import Pt, RGBColor

from matcher import MatchResult
from date_detector import MURLI_TYPE_SAKAR, MURLI_TYPE_AVYAKT


_UNSAFE_FILENAME_CHARS = re.compile(r'[\\/:*?"<>|]')


def _sanitize_for_filename(text: str, max_len: int = 40) -> str:
    text = unicodedata.normalize("NFC", text)
    text = re.sub(r'\s+', '_', text.strip())
    text = _UNSAFE_FILENAME_CHARS.sub('', text)  # only drop chars illegal in filenames
    return text[:max_len] if text else "search"


def build_output_filename(search_text: str, murli_type: str) -> str:
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    label = murli_type if murli_type in (MURLI_TYPE_SAKAR, MURLI_TYPE_AVYAKT) else "Murli"
    safe_text = _sanitize_for_filename(search_text)
    return f"{safe_text}_{label}_{ts}.docx"


def write_results_docx(
    matches: List[MatchResult],
    search_text: str,
    murli_type: str,
    source_file_name: str,
    output_dir: str,
) -> str:
    os.makedirs(output_dir, exist_ok=True)
    filename = build_output_filename(search_text, murli_type)
    output_path = os.path.join(output_dir, filename)

    document = docx.Document()

    title = document.add_heading(f'Search Results: "{search_text}"', level=1)

    meta = document.add_paragraph()
    meta.add_run(f"Source file: ").bold = True
    meta.add_run(source_file_name)
    meta.add_run(f"\nMurli type: ").bold = True
    meta.add_run(murli_type)
    meta.add_run(f"\nTotal matches found: ").bold = True
    meta.add_run(str(len(matches)))
    document.add_paragraph()  # spacer

    if not matches:
        document.add_paragraph("No matches were found for the given text.")
    else:
        for i, m in enumerate(matches, start=1):
            heading = document.add_paragraph()
            run = heading.add_run(f"Result {i}")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

            info = document.add_paragraph()
            info.add_run("Date: ").bold = True
            info.add_run(m.date_string)

            excerpt = document.add_paragraph()
            before = m.context_text[: m.bold_start]
            matched = m.context_text[m.bold_start : m.bold_end]
            after = m.context_text[m.bold_end :]

            if before:
                excerpt.add_run(before)
            bold_run = excerpt.add_run(matched)
            bold_run.bold = True
            if after:
                excerpt.add_run(after)

            document.add_paragraph()  # spacer between results
    document.styles['Normal'].font.size = Pt(20)
    document.save(output_path)
    return output_path
